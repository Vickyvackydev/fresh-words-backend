import sys
import os
import re
import docx
from docx.shared import Pt

def convert(txt_path, docx_path):
    doc = docx.Document()
    
    # Set default style to Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
        lines = f.readlines()

    # Clean lines and filter out empty strings
    cleaned_lines = [line.strip() for line in lines if line.strip()]

    months = ['JANUARY', 'FEBRUARY', 'MARCH', 'APRIL', 'MAY', 'JUNE', 'JULY', 'AUGUST', 'SEPTEMBER', 'OCTOBER', 'NOVEMBER', 'DECEMBER', 'JAN', 'FEB', 'MAR', 'APR', 'JUN', 'JUL', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC']
    months_pattern = r'^(' + '|'.join(months) + r')\s+\d{1,2}\b'

    section_patterns = [
        r'^PRAYER\b', r'^PRAYER POINT\b', r'^REFLECTION\b', r'^CONFESSION\b',
        r'^FURTHER READING\b', r'^MEMORY VERSE\b', r'^KEY POINT\b', r'^ACTION POINT\b',
        r'^DECLARATION\b'
    ]

    paragraphs = []
    current_p_words = []
    p_styles = [] # List of tuples: (is_bold, font_size)

    state = "normal" # "date", "title", "normal"

    for i, line in enumerate(cleaned_lines):
        next_line = cleaned_lines[i+1] if i+1 < len(cleaned_lines) else None
        
        current_p_words.append(line)
        
        # Heuristics to determine if we should split/end the current paragraph
        split = False
        is_bold = False
        font_size = 11
        
        # Rule 1: Next line is a Date -> End current paragraph
        if next_line and re.match(months_pattern, next_line, re.IGNORECASE) and len(next_line) < 20:
            split = True
            
        # Rule 2: Current line is a Date -> End current paragraph and make it a bold date heading
        elif re.match(months_pattern, line, re.IGNORECASE) and len(line) < 20:
            split = True
            is_bold = True
            font_size = 14
            state = "title"
            
        # Rule 3: Current line is the Title (the line immediately after a Date) -> Bold title heading
        elif state == "title":
            split = True
            is_bold = True
            font_size = 16
            state = "normal"
            
        # Rule 4: Next line is a Section Header -> End current paragraph
        elif next_line and any(re.match(pat, next_line, re.IGNORECASE) for pat in section_patterns) and len(next_line) < 35:
            split = True
            
        # Rule 5: Current line is a Section Header -> End current paragraph and bold it
        elif any(re.match(pat, line, re.IGNORECASE) for pat in section_patterns) and len(line) < 35:
            split = True
            is_bold = True
            font_size = 12
            
        # Rule 6: Current line ends with sentence ending punctuation and is short (< 55 chars)
        elif line[-1] in ['.', '?', '!', '"', ')'] and len(line) < 55:
            split = True
            
        # Rule 7: Current line is extremely short (< 30 chars), likely a stray tag or final line
        elif len(line) < 30:
            split = True
            
        if split:
            merged_text = " ".join(current_p_words).strip()
            merged_text = re.sub(r'\s+', ' ', merged_text)
            if merged_text:
                paragraphs.append((merged_text, is_bold, font_size))
            current_p_words = []

    if current_p_words:
        merged_text = " ".join(current_p_words).strip()
        merged_text = re.sub(r'\s+', ' ', merged_text)
        if merged_text:
            paragraphs.append((merged_text, False, 11))

    # Save to Word Document
    for text, is_bold, font_size in paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = is_bold
        p.style.font.size = Pt(font_size)

    doc.save(docx_path)
    print(f"Successfully converted TXT to DOCX: {docx_path}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) >= 3:
        txt_path = sys.argv[1]
        docx_path = sys.argv[2]
    else:
        txt_path = r"C:\Users\USER\Downloads\Holiness_Devotional_Extracted_New.txt"
        docx_path = r"C:\Users\USER\Downloads\Holiness_Devotional_Extracted_New.docx"
    convert(txt_path, docx_path)
